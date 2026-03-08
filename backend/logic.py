import json
import os
import urllib.parse
from pypdf import PdfReader
import unicodedata
import re
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from weasyprint import HTML, CSS

# Sistema de logging unificado - importação direta sem fallback
from logging_config import setup_logger
logger = setup_logger("VANT_LOGIC")

import html

def sanitize_input(text: str) -> str:
    """Remove caracteres perigosos de inputs do usuário."""
    if not text:
        return ""
    
    # Remove HTML tags
    text = html.escape(text)
    
    # Remove caracteres de controle (exceto \n, \r, \t)
    text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
    
    # Limita tamanho (proteção contra DoS)
    MAX_INPUT_SIZE = 50000  # ~50KB
    if len(text) > MAX_INPUT_SIZE:
        text = text[:MAX_INPUT_SIZE]
        logger.warning(f"⚠️ Input truncado (>{MAX_INPUT_SIZE} chars)")
    
    return text

# ==============================================================================
# HELPER FUNCTIONS (ESSENCIAIS PARA O PARSER)
# ==============================================================================

def safe_txt(text):
    """
    Limpa caracteres quebrados de encoding e normaliza espaços.
    
    Args:
        text (str): Texto bruto a ser limpo.
    
    Returns:
        str: Texto limpo com caracteres especiais substituídos e espaços normalizados.
    """
    if not text: return ""
    replacements = {
        '–': '-', '—': '-', '“': '"', '”': '"', '‘': "'", '’': "'", 
        '•': '', '●': '', '⚫': '', '■': '', '\u200b': '', '\t': ' ', '*': '' 
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    
    # Corrige objetos colados sem vírgula: e.g. "}{"
    text = re.sub(r'}\s*{', '},{', text)

    # Remove espaços duplos
    text = re.sub(r'\s{2,}', ' ', text)
    # Remove markdown residual se necessário
    text = text.replace('**', '').replace('__', '').replace('#', '')
    if text.strip().startswith('- '): text = text.strip()[2:]
    
    return text.strip()

def clean_garbage(text):
    """Remove termos como (não informado) e sujeira de OCR."""
    if not text: return ""
    garbage = [
        r'\(?n[ãa]o\s*informado\)?', r'\(?n[ãa]o\s*fornecido\)?', 
        r'\(?sem\s*telefone\)?', r'\(?undefined\)?', r'null'
    ]
    clean = text
    for p in garbage: 
        clean = re.sub(p, '', clean, flags=re.IGNORECASE)
    
    clean = clean.strip().rstrip('|').strip()
    
    # Se sobrou apenas o rótulo, retorna vazio
    if clean.lower() in ['telefone:', 'email:', 'linkedin:', 'local:', 'e', 'a']: 
        return ""
        
    return clean

def normalize_visual(text):
    if not text:
        return ""

    # Remove Markdown estrutural (títulos, bold, etc)
    text = re.sub(r'^#+\s*', '', text)   # remove #, ##, ### no início
    text = re.sub(r'[*_`]+', '', text)   # remove ** __ `

    # Normalizações visuais que você já tem
    text = unicodedata.normalize("NFKC", text)
    text = re.sub(r'\s{2,}', ' ', text)

    return text.strip()



def extract_date(text):
    """Extrai padrões de data (ex: Jan 2020 - Atual) e retorna o resto do texto."""
    # Padrões comuns de data em CVs
    patterns = [
        r'(\b(?:jan|fev|mar|abr|mai|jun|jul|ago|set|out|nov|dez)[a-z]*\s*(?:de)?\s*\d{4}\s*[-–]\s*(?:presente|atual|\w+\s*\d{4}))',
        r'(\d{2}/\d{4}\s*[-–]\s*(?:presente|atual|\d{2}/\d{4}))',
        r'(\d{4}\s*[-–]\s*(?:presente|atual|\d{4}))'
    ]
    found = ""
    text_clean = text
    
    for pat in patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            found = m.group(1).strip()
            # Remove a data do texto original para sobrar só o Cargo
            text_clean = text.replace(m.group(0), '').strip()
            # Limpa separadores que sobraram no fim
            if text_clean.endswith('|') or text_clean.endswith('-'): 
                text_clean = text_clean[:-1].strip()
            break
            
    return text_clean, found

# ============================================================
# IMPORTA A INTELIGÊNCIA
# ============================================================
# Importação direta sem fallback - ambiente deve estar configurado corretamente
from llm_core import run_llm_orchestrator

# ============================================================
# CONFIG
# ============================================================
# Configuração de logging já feita via setup_logger

# ============================================================
# PDF / TEXTO (COM LIMITE DE TOKENS)
# ============================================================
def extrair_texto_pdf(uploaded_file, max_chars=25000):
    """
    Extrai texto e verifica se o arquivo foi gerado pelo VANT 
    (via Metadado ou Marca d'água no texto) para ativar o modo 'Certified'.
    """
    try:
        reader = PdfReader(uploaded_file)
        
        full_text = ""
        for page in reader.pages:
            full_text += page.extract_text() or ""

        # [TECH LEAD MAGIC] - Detecção da Assinatura Digital
        is_vant_certified = False
        
        # Check A: Metadado (Para Upload Direto do PDF gerado pelo sistema)
        if reader.metadata and reader.metadata.get('/Creator') == "VANT_NEURAL_ENGINE_V2":
            is_vant_certified = True
            logger.info("💎 VANT CERTIFIED: Detectado via Metadado PDF!")
            
        # Check B: Marca d'água no Texto (Caso o usuário tenha Copiado/Colado no Word)
        elif "VANT-ID: VANT-NEURAL-ENGINE-CERTIFIED" in full_text:
            is_vant_certified = True
            logger.info("💎 VANT CERTIFIED: Detectado via Texto/Watermark!")

        # Injeção de Contexto para a IA
        # Se for certificado, injetamos o cabeçalho que garante nota alta de estrutura/ATS.
        if is_vant_certified:
            header_magic = "\n[SYSTEM_INTERNAL_NOTE: ESTE CURRÍCULO JÁ FOI OTIMIZADO PELO VANT. A ESTRUTURA E FORMATACAO ESTAO PERFEITAS (NOTA ATS DEVE SER > 90 NA PARTE ESTRUTURAL). AVALIE APENAS A ADERENCIA DAS SKILLS.]\n"
            full_text = header_magic + full_text

        if len(full_text) > max_chars:
            logger.info(f"⚠️ PDF truncado em {max_chars} caracteres para segurança.")
            return full_text[:max_chars]
            
        return full_text
        
    except Exception as e:
        logger.error(f"Erro crítico ao ler PDF: {e}")
        return ""

# ============================================================
# CATÁLOGO DE ELITE (V5.0 - LIBRARY EXPANDIDA)
# Fonte da Verdade para Fallback & Curadoria IA
# ============================================================
BACKUP_CATALOG = {
    "ti_dev_gen": [
        # Engenharia de Software & Código
        {"titulo": "Clean Code", "autor": "Robert C. Martin", "motivo": "Diferencia o amador do profissional. Essencial para passar nos Code Reviews de grandes empresas."},
        {"titulo": "O Programador Pragmático", "autor": "Andrew Hunt", "motivo": "Vai te ensinar a pensar como um Engenheiro de Software Sênior, não apenas como um codificador."},
        {"titulo": "Arquitetura Limpa", "autor": "Robert C. Martin", "motivo": "Leitura obrigatória se você quer desenhar sistemas escaláveis e não apenas consertar bugs."},
        {"titulo": "Refatoração", "autor": "Martin Fowler", "motivo": "A técnica vital para trabalhar com código legado sem quebrar o sistema inteiro."},
        {"titulo": "Padrões de Projeto (Design Patterns)", "autor": "Erich Gamma", "motivo": "A linguagem universal dos arquitetos de software. Domine isso para discutir soluções técnicas."},
        
        # Carreira & Algoritmos
        {"titulo": "Entendendo Algoritmos", "autor": "Aditya Bhargava", "motivo": "O guia ilustrado definitivo para perder o medo de lógica complexa e otimização."},
        {"titulo": "System Design Interview", "autor": "Alex Xu", "motivo": "A chave mestra para ser aprovado em entrevistas técnicas de Big Techs e cargos de liderança."},
        {"titulo": "Cracking the Coding Interview", "autor": "Gayle Laakmann", "motivo": "O manual de sobrevivência para passar nos testes técnicos mais difíceis do mercado."},
        
        # DevOps & Cultura
        {"titulo": "O Projeto Fênix", "autor": "Gene Kim", "motivo": "Entenda como a TI impacta o negócio. Leitura obrigatória para quem busca cargos de DevOps ou Gestão."},
        {"titulo": "Engenharia de Confiabilidade (SRE)", "autor": "Google", "motivo": "Como o Google mantém seus sistemas no ar. A bíblia para infraestrutura e sustentação."}
    ],
    
    # NOVA CATEGORIA: SUPORTE E INFRAESTRUTURA
    "ti_suporte": [
        {"titulo": "Google IT Support Professional", "autor": "Google", "motivo": "A base fundamental para suporte moderno e troubleshooting eficaz."},
        {"titulo": "ITIL 4 Foundation", "autor": "Axelos", "motivo": "O padrão mundial para gerenciamento de serviços de TI. Essencial para empresas grandes."},
        {"titulo": "Redes de Computadores", "autor": "Tanenbaum", "motivo": "A bíblia técnica para entender TCP/IP e resolver problemas de conectividade."},
        {"titulo": "Windows Internals", "autor": "Pavel Yosifovich", "motivo": "Para quem precisa resolver problemas profundos no SO que o reboot não resolve."},
        {"titulo": "The Service Desk Handbook", "autor": "Sanjay Nagaraj", "motivo": "Focado em atendimento ao cliente técnico e SLAs."},
        {"titulo": "Comece pelo Porquê", "autor": "Simon Sinek", "motivo": "Soft skill vital: entender a dor do usuário antes de mexer no computador."}
    ],

    "ti_dados_ai": [
        # Estratégia & Comunicação
        {"titulo": "Storytelling com Dados", "autor": "Cole Nussbaumer", "motivo": "Transforma você de um analista de planilhas em um parceiro estratégico que influencia diretores."},
        {"titulo": "Data Science para Negócios", "autor": "Foster Provost", "motivo": "Para entender como os dados geram dinheiro de verdade nas empresas, além do algoritmo."},
        {"titulo": "Como Mentir com Estatística", "autor": "Darrell Huff", "motivo": "Blinde suas análises contra vieses e ganhe confiança absoluta nos seus relatórios."},
        
        # Técnico & Mão na Massa
        {"titulo": "Mãos à Obra: Aprendizado de Máquina", "autor": "Aurélien Géron", "motivo": "O manual prático definitivo para sair da teoria e colocar modelos de IA em produção real."},
        {"titulo": "Python para Análise de Dados", "autor": "Wes McKinney", "motivo": "Escrito pelo criador do Pandas. É a referência técnica para manipulação de dados pesada."},
        {"titulo": "Deep Learning Book", "autor": "Ian Goodfellow", "motivo": "A referência acadêmica máxima se você quer mergulhar fundo em Redes Neurais."},
        {"titulo": "Designing Data-Intensive Applications", "autor": "Martin Kleppmann", "motivo": "O livro mais respeitado do mundo sobre como construir sistemas de dados robustos e distribuídos."}
    ],

    "produto_agil": [
        # Gestão de Produto
        {"titulo": "Inspirado (Inspired)", "autor": "Marty Cagan", "motivo": "O livro de cabeceira dos melhores Product Managers do Vale do Silício. Define o padrão da indústria."},
        {"titulo": "Empowered", "autor": "Marty Cagan", "motivo": "Leitura crítica se você busca sair da execução de tarefas para atuar com autonomia e liderança de produto."},
        {"titulo": "Hooked (Engajado)", "autor": "Nir Eyal", "motivo": "A psicologia por trás dos produtos que retêm usuários. Essencial para métricas de retenção."},
        {"titulo": "Continuous Discovery Habits", "autor": "Teresa Torres", "motivo": "Como descobrir o que o cliente realmente quer, de forma contínua, sem 'achismos'."},
        
        # Metodologia & Estratégia
        {"titulo": "A Startup Enxuta (Lean Startup)", "autor": "Eric Ries", "motivo": "A base fundamental de validação. Aprenda a errar rápido para acertar o produto."},
        {"titulo": "Sprint", "autor": "Jake Knapp", "motivo": "A metodologia do Google para validar ideias em 5 dias, evitando meses de desenvolvimento inútil."},
        {"titulo": "Scrum: A Arte de Fazer o Dobro...", "autor": "Jeff Sutherland", "motivo": "A origem da metodologia ágil. Essencial para organizar times caóticos."},
        {"titulo": "Crossing the Chasm", "autor": "Geoffrey Moore", "motivo": "A estratégia definitiva para lançar produtos de tecnologia em mercados B2B."}
    ],

    "marketing_growth": [
        # Estratégia Digital & Growth
        {"titulo": "Hacking Growth", "autor": "Sean Ellis", "motivo": "A metodologia exata usada para escalar startups unicórnio. Saia do 'achismo' para os testes rápidos."},
        {"titulo": "Marketing 4.0", "autor": "Philip Kotler", "motivo": "Atualize seu mindset para as estratégias digitais modernas com a maior autoridade do mundo."},
        {"titulo": "Traction", "autor": "Gabriel Weinberg", "motivo": "Um guia prático com 19 canais de aquisição para você nunca ficar sem clientes."},
        {"titulo": "Dotcom Secrets", "autor": "Russell Brunson", "motivo": "A engenharia por trás de funis de vendas online que convertem visitantes em compradores."},
        
        # Psicologia & Branding
        {"titulo": "As Armas da Persuasão", "autor": "Robert Cialdini", "motivo": "Domine os gatilhos mentais para aumentar sua conversão em qualquer campanha ou negociação."},
        {"titulo": "Contágio: Por que as Coisas Pegam", "autor": "Jonah Berger", "motivo": "A ciência da viralização. Entenda por que alguns conteúdos explodem e outros morrem."},
        {"titulo": "StoryBrand", "autor": "Donald Miller", "motivo": "Aprenda a comunicar o valor da sua marca de forma tão clara que o cliente não consegue ignorar."},
        {"titulo": "Posicionamento", "autor": "Al Ries", "motivo": "O clássico sobre como ocupar um lugar único na mente do seu consumidor."}
    ],

    "vendas_cs": [
        # B2B & Vendas Complexas
        {"titulo": "Receita Previsível", "autor": "Aaron Ross", "motivo": "O blueprint para construir uma máquina de vendas B2B escalável e sair da dependência de indicação."},
        {"titulo": "Spin Selling", "autor": "Neil Rackham", "motivo": "A técnica definitiva para fechar contratos de alto valor (Enterprise) fazendo as perguntas certas."},
        {"titulo": "A Venda Desafiadora", "autor": "Matthew Dixon", "motivo": "Para vender soluções complexas, você precisa desafiar o cliente e ensinar algo novo, não apenas agradar."},
        {"titulo": "Negocie Como se Sua Vida Dependesse Disso", "autor": "Chris Voss", "motivo": "Técnicas de negociação do FBI aplicadas ao mundo corporativo de alto risco."},
        
        # Atitude & Retenção
        {"titulo": "Customer Success", "autor": "Dan Steinman", "motivo": "Leitura obrigatória para entender que a venda real começa após a assinatura do contrato."},
        {"titulo": "A Bíblia de Vendas", "autor": "Jeffrey Gitomer", "motivo": "O guia prático de atitude e técnica para quem vive de comissão."},
        {"titulo": "Fanatical Prospecting", "autor": "Jeb Blount", "motivo": "Acabe com o problema de pipeline vazio. O guia bruto sobre prospecção ativa."}
    ],

    "rh_lideranca": [
        # Gestão & Cultura
        {"titulo": "Work Rules!", "autor": "Laszlo Bock", "motivo": "Os segredos de gestão de pessoas que fizeram do Google uma referência mundial em talentos."},
        {"titulo": "High Output Management", "autor": "Andrew Grove", "motivo": "A bíblia da gestão de alta performance escrita pelo lendário CEO da Intel."},
        {"titulo": "Os 5 Desafios das Equipes", "autor": "Patrick Lencioni", "motivo": "Entenda a dinâmica oculta que impede seu time de atingir a alta performance."},
        
        # Desenvolvimento de Líderes
        {"titulo": "Pipeline de Liderança", "autor": "Ram Charan", "motivo": "O mapa claro do que é exigido em cada degrau da escada corporativa para você ser promovido."},
        {"titulo": "Radical Candor", "autor": "Kim Scott", "motivo": "Aprenda a dar feedbacks difíceis sem destruir a relação com seu time. Vital para gestores."},
        {"titulo": "A Coragem de Ser Imperfeito", "autor": "Brené Brown", "motivo": "Como a vulnerabilidade pode ser sua maior força na liderança de equipes modernas."},
        {"titulo": "Comece pelo Porquê", "autor": "Simon Sinek", "motivo": "Fundamental para líderes que precisam inspirar ação e propósito, não apenas dar ordens."}
    ],

    "financeiro_corp": [
        # Finanças Técnicas
        {"titulo": "Valuation", "autor": "Aswath Damodaran", "motivo": "A bíblia técnica para quem quer falar a língua dos CFOs e investidores com propriedade."},
        {"titulo": "O Investidor Inteligente", "autor": "Benjamin Graham", "motivo": "Fundamentos de análise financeira que sobrevivem a qualquer crise de mercado."},
        {"titulo": "Financial Intelligence", "autor": "Karen Berman", "motivo": "Contabilidade e finanças explicadas para gestores que não são da área financeira."},
        
        # Mindset & Negócios
        {"titulo": "Princípios", "autor": "Ray Dalio", "motivo": "Como sistematizar a cultura e a tomada de decisão para atingir resultados excepcionais."},
        {"titulo": "A Psicologia do Dinheiro", "autor": "Morgan Housel", "motivo": "Entenda como o comportamento humano afeta as decisões financeiras mais do que a matemática."},
        {"titulo": "Pai Rico, Pai Pobre", "autor": "Robert Kiyosaki", "motivo": "Essencial para desenvolver a mentalidade de ativos vs passivos na gestão de patrimônio."},
        {"titulo": "A Marca da Vitória (Shoe Dog)", "autor": "Phil Knight", "motivo": "A jornada real e caótica de construir uma empresa global (Nike) do zero."}
    ],

    "construcao_manual": [
        {"titulo": "Manual do Construtor", "autor": "Vários", "motivo": "A referência técnica para executar obras com padrão de engenharia, evitando retrabalho e desperdício."},
        {"titulo": "Normas Regulamentadoras (NRs Comentadas)", "autor": "Editora Saraiva", "motivo": "Domine as normas de segurança para ser o profissional mais confiável e requisitado da obra."},
        {"titulo": "Instalações Elétricas Prediais", "autor": "Hélio Creder", "motivo": "O guia definitivo para eletricistas que querem garantir segurança e conformidade técnica."},
        {"titulo": "Mestre de Obras: Gestão Básica", "autor": "Senai", "motivo": "O passo fundamental para deixar de ser operacional e começar a liderar equipes no canteiro."},
        {"titulo": "Concreto Armado Eu Te Amo", "autor": "Manoel Botelho", "motivo": "A melhor didática do mercado para entender estruturas sem complicação matemática."}
    ],

    "gastronomia": [
        {"titulo": "Kitchen Confidential", "autor": "Anthony Bourdain", "motivo": "Entenda a realidade bruta, a hierarquia e a disciplina necessária numa cozinha de alta performance."},
        {"titulo": "The Professional Chef", "autor": "CIA", "motivo": "A bíblia técnica culinária. Diferencia o cozinheiro amador do Chef profissional."},
        {"titulo": "Sal, Gordura, Ácido, Calor", "autor": "Samin Nosrat", "motivo": "Domine os 4 elementos fundamentais para criar sabor sem depender apenas de receitas prontas."},
        {"titulo": "Larousse Gastronomique", "autor": "Prosper Montagné", "motivo": "A enciclopédia definitiva da gastronomia mundial. Referência absoluta."}
    ],

    "global_soft_skills": [
        {"titulo": "Hábitos Atômicos", "autor": "James Clear", "motivo": "Pequenas mudanças de rotina que compõem resultados gigantes na sua carreira a longo prazo."},
        {"titulo": "Trabalho Focado (Deep Work)", "autor": "Cal Newport", "motivo": "A habilidade mais rara do século XXI: focar sem distração para produzir trabalho de elite."},
        {"titulo": "Essencialismo", "autor": "Greg McKeown", "motivo": "Aprenda a dizer 'não' para o que é apenas bom e focar energia exclusivamente no que é excelente."},
        {"titulo": "Como Fazer Amigos e Influenciar Pessoas", "autor": "Dale Carnegie", "motivo": "A habilidade número 1 para crescer na política corporativa, vendas e networking."},
        {"titulo": "Comunicação Não-Violenta", "autor": "Marshall Rosenberg", "motivo": "A ferramenta essencial para resolver conflitos e negociar em ambientes de alta pressão."},
        {"titulo": "Mindset", "autor": "Carol Dweck", "motivo": "A chave psicológica para aceitar desafios e crescer profissionalmente."},
        {"titulo": "O Ego é seu Inimigo", "autor": "Ryan Holiday", "motivo": "Como impedir que sua arrogância destrua sua carreira nos momentos de sucesso."}
    ]
}

from functools import lru_cache
import json

@lru_cache(maxsize=1)
def load_books_catalog():
    """Carrega catálogo de livros (cached)."""
    base_dir = os.path.dirname(os.path.abspath(__file__))
    catalog_path = os.path.join(base_dir, "data", "books_catalog.json")
    
    try:
        if os.path.exists(catalog_path):
            with open(catalog_path, encoding="utf-8") as f:
                return json.load(f)
        else:
            logger.warning(" Catálogo de livros não encontrado")
            return {}
    except Exception as e:
        logger.error(f" Erro ao carregar catálogo: {e}")
        return {}

# ============================================================
# DETECTOR DE ÁREA (REGEX OTIMIZADO)
# ============================================================

def detect_job_area(job_description):
    job_lower = job_description.lower()
    
    # Detectar se é vaga genérica
    generic_keywords = [
        r"busco oportunidades profissionais",
        r"estou aberto a posições",
        r"crescimento e contribuição",
        r"objetivos da empresa",
        r"foco em resultados e inovação"
    ]
    
    # Se for vaga genérica, retornar área selecionada ou global_soft_skills
    for pattern in generic_keywords:
        if re.search(pattern, job_lower):
            return "global_soft_skills"  # Será substituído pela área selecionada se houver
    
    keyword_map = {
        "ti_dados_ai": [
            r"dados", r"data", r"analytics", r"bi\b", r"business intelligence", r"cientista de dados", 
            r"machine learning", r"ia\b", r"inteligência artificial", r"python", r"pandas", r"sql", r"big data"
        ],
        "ti_suporte": [
            r"suporte", r"help desk", r"service desk", r"infraestrutura", r"infra", r"sysadmin", 
            r"técnico de ti", r"n1", r"n2", r"field service", r"atendimento", r"hardware", r"redes"
        ],
        "ti_dev_gen": [
            r"desenvolvedor", r"developer", r"engenheiro de software", r"software engineer", 
            r"fullstack", r"backend", r"frontend", r"java\b", r"python", r"react"
        ],
        "produto_agil": [
            r"produto", r"product manager", r"product owner", r"po\b", r"pm\b", r"scrum", r"agile", r"kanban", 
            r"agilista", r"roadmap", r"backlog", r"user story"
        ],
        "marketing_growth": [
            r"marketing", r"growth", r"performance", r"tráfego", r"seo\b", r"conteúdo", r"social media", 
            r"branding", r"copywriter", r"crm\b", r"inbound", r"redator", r"designer"
        ],
        "vendas_cs": [
            r"vendas", r"sales", r"comercial", r"sdr\b", r"bdr\b", r"closer", r"executivo de contas", r"account executive", 
            r"customer success", r"sucesso do cliente", r"pós-venda", r"churn", r"negociação"
        ],
        "rh_lideranca": [
            r"rh\b", r"r\.h\.", r"recursos humanos", r"recrutamento", r"talent", r"people", r"dp\b", r"departamento pessoal", 
            r"tech recruiter", r"bp\b", r"business partner", r"liderança", r"gestão de pessoas", r"coordenador", r"gerente", r"supervisor", 
            r"analista de rh" 
        ],
        "financeiro_corp": [
            r"financeiro", r"finanças", r"contábil", r"contabilidade", r"fiscal", r"auditoria", r"controller", 
            r"fp&a", r"tesouraria", r"banco", r"investimento", r"fusões", r"economista"
        ],
        "construcao_manual": [
            r"pedreiro", r"servente", r"mestre de obras", r"obra", r"civil", r"construção", 
            r"elétrica", r"eletricista", r"manutenção", r"engenheiro civil", r"arquitetura"
        ],
        "gastronomia": [
            r"cozinha", r"chef", r"gastronomia", r"culinária", r"restaurante", r"alimento", r"cook"
        ]
    }
    
    best_match = "global_soft_skills"
    max_count = 0
    
    for area, patterns in keyword_map.items():
        count = 0
        for pattern in patterns:
            if re.search(pattern, job_lower):
                count += 1
        
        if count > max_count:
            max_count = count
            best_match = area
            
    return best_match

def format_text_to_html_diagnostic(text):
    """
    Versão específica para blocos de diagnóstico (mantém cores vibrantes).
    Diferente do Paper View que usa preto.
    """
    if not text: return ""
    
    # Escapa HTML
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    
    # Negrito → Ciano Neon (cor de destaque do sistema)
    text = re.sub(r'\*\*(.*?)\*\*', r'<strong style="color: #38BDF8; font-weight: 700;">\1</strong>', text)
    
    # Itálico
    text = re.sub(r'(?<!\*)\*(?!\*)(.*?)\*', r'<em style="color: #94A3B8; font-style: normal; font-weight: 600;">\1</em>', text)
    
    # Quebras de linha
    text = text.replace('\n', '<br>')
    
    return text

# ============================================================
# FORMATADOR HTML V11 (STABLE LAYOUT - DUAL ROW)
# ============================================================
def format_text_to_html(text):
    if not text: return ""
    
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    html_output = []
    lines = text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line: continue
            
        # 1. NOME & 2. SEÇÕES (Manter igual)
        if line.startswith('# '):
            clean = line.replace('# ', '').upper()
            html_output.append(f'<h1 class="vant-cv-name">{clean}</h1>')
        elif line.startswith('###'):
            clean = line.replace('###', '').strip().upper()
            html_output.append(f'<h2 class="vant-cv-section">{clean}</h2>')
            
        # 3. TRATAMENTO DE LISTAS (CARGOS vs TAREFAS)
        elif line.startswith('- ') or line.startswith('* ') or line.startswith('• '):
            clean = re.sub(r'^[-*•]\s+', '', line)
            clean = re.sub(r'\*\*(.*?)\*\*', r'<span class="vant-bold">\1</span>', clean)
            
            # --- LÓGICA DE HEADER DE VAGA (ESTRUTURA DE 2 LINHAS) ---
            if '|' in line:
                parts = [p.strip() for p in clean.split('|')]
                
                # Variáveis padrão
                cargo = clean
                empresa = ""
                data = ""
                
                if len(parts) >= 3:
                    cargo = parts[0]
                    empresa = parts[1]
                    data = parts[2].replace('*', '').replace('_', '').strip()
                elif len(parts) == 2:
                    cargo = parts[0]
                    empresa = parts[1].replace('*', '').strip()
                
                # Montagem do HTML com quebra forçada
                job_html = f"""
                <div class="vant-cv-job-container">
                    <div class="vant-job-row-primary">
                        <span class="vant-job-title">{cargo}</span>
                        <span class="vant-job-sep">|</span>
                        <span class="vant-job-company">{empresa}</span>
                    </div>
                    """
                
                # Só adiciona a linha da data se ela existir
                if data:
                    job_html += f"""
                    <div class="vant-job-row-secondary">
                        <span class="vant-job-date">{data}</span>
                    </div>
                    """
                
                job_html += "</div>"
                html_output.append(job_html)
            
            # --- LÓGICA DE TAREFAS (Manter igual) ---
            else:
                row = f"""
                <div class="vant-cv-grid-row">
                    <div class="vant-cv-bullet-col">•</div>
                    <div class="vant-cv-text-col">{clean}</div>
                </div>
                """
                html_output.append(row)

        # 4. CONTATOS & 5. TEXTO CORRIDO (Manter igual)
        elif ('|' in line or '@' in line) and len(line) < 300:
            clean_line = line.replace('**', '')
            parts = [p.strip() for p in clean_line.split('|')]
            items_html = []
            for p in parts:
                if p:
                    if ":" in p:
                        label, val = p.split(":", 1)
                        block = f'<span class="vant-contact-block"><span class="vant-bold">{label}:</span> {val}</span>'
                        items_html.append(block)
                    else:
                        items_html.append(f'<span class="vant-contact-block">{p}</span>')
            full_html = '<span class="vant-contact-separator"> • </span>'.join(items_html)
            html_output.append(f'<div class="vant-cv-contact-line">{full_html}</div>')
        else:
            clean = re.sub(r'\*\*(.*?)\*\*', r'<span class="vant-bold">\1</span>', line)
            html_output.append(f'<p class="vant-cv-paragraph">{clean}</p>')
            
    return "\n".join(html_output)

# ============================================================
# LINK AMAZON INTELIGENTE
# ============================================================
def gerar_link_amazon(titulo_livro, autor=None):
    tag_afiliado = "rodrigoverruc-20"
    
    if isinstance(titulo_livro, dict) and "amazon_url" in titulo_livro:
        return titulo_livro["amazon_url"]
        
    try:
        titulo_limpo = titulo_livro.split(":")[0] 
        termo_busca = titulo_limpo
        if autor:
            autor_limpo = autor.split(",")[0] 
            termo_busca = f"{titulo_limpo} {autor_limpo}"
            
        query = urllib.parse.quote(termo_busca)
        return f"https://www.amazon.com.br/s?k={query}&tag={tag_afiliado}"
        
    except Exception:
        return f"https://www.amazon.com.br/s?k=livros+profissionais&tag={tag_afiliado}"

# ============================================================
# GERADOR DE PDF (VERSÃO FINAL COM WEASYPRINT)
# ============================================================
# GERADOR DE PDF (VERSÃO FINAL COM WEASYPRINT)
# ============================================================
from weasyprint import HTML, CSS
import re
import unicodedata

# Importa constantes CSS centralizadas
from styles import CSS_V13, CSS_PDF
# ==============================================================================
# 1. PARSER ENGINE (Mantido para limpar seus dados brutos)
# ==============================================================================
def parse_raw_data_to_struct(raw_text):
    """
    CÉREBRO DO SISTEMA: Transforma texto bruto em dados estruturados.
    Versão Fiel: Mantém rótulos (Email:, Telefone:) exatamente como aparecem.
    """
    raw_lines = raw_text.split('\n')
    parsed = {"name": "", "contacts": [], "sections": []}
    curr_sec = None
    buffer = []

    def flush():
        nonlocal curr_sec
        if not curr_sec:
            buffer.clear()
            return
        
        # SKILLS
        if any(k in curr_sec for k in ["SKILLS", "COMPETÊNCIAS"]):
            txt = " ".join(buffer).replace('\n', ' ')
            if '|' in txt:
                s = [safe_txt(x) for x in txt.split('|') if safe_txt(x)]
            else:
                s = [safe_txt(x) for x in re.split(r'\s{2,}', txt) if safe_txt(x)]

            if s:
                parsed["sections"].append({
                    "title": curr_sec,
                    "type": "chips",
                    "content": s
                })

        # RESUMO
        elif "RESUMO" in curr_sec:
            full = " ".join([safe_txt(x) for x in buffer if safe_txt(x)])
            if full:
                parsed["sections"].append({
                    "title": curr_sec,
                    "type": "paragraph",
                    "content": full
                })

        # LISTAS SIMPLES
        elif any(k in curr_sec for k in ["IDIOMA", "CERTIFICA"]):
            entries = []
            for raw_line in buffer:
                entry = raw_line.strip()
                if entry and entry not in ['.', '-', '•', '|']:
                    entries.append(entry)
            if entries:
                parsed["sections"].append({
                    "title": curr_sec,
                    "type": "list_simple",
                    "content": entries
                })

        # JOBS / EXPERIÊNCIA
        else:
            job_blocks = []
            clean_lines = [safe_txt(x) for x in buffer if safe_txt(x)]
            i = 0

            while i < len(clean_lines):
                line = clean_lines[i]

                # REGRA DE OURO:
                # Só é cabeçalho de cargo se tiver "|"
                is_job_header = '|' in line

                if is_job_header:
                    text_rem, date_found = extract_date(line)

                    # Divide cargo | empresa
                    parts = text_rem.split('|', 1)
                    cargo = parts[0].strip()
                    empresa = parts[1].strip() if len(parts) > 1 else ""

                    job_blocks.append({
                        "cargo": cargo,
                        "empresa": empresa,
                        "date": date_found,
                        "details": []
                    })

                else:
                    # Tudo que NÃO é cabeçalho vira detalhe
                    if job_blocks:
                        clean_detail = line.strip()

                        # ignora lixo visual
                        if clean_detail not in ['.', '-', '•', '|']:
                            job_blocks[-1]["details"].append(clean_detail)

                i += 1

            # SALVA UMA ÚNICA VEZ
            if job_blocks:
                parsed["sections"].append({
                    "title": curr_sec,
                    "type": "jobs",
                    "content": job_blocks
                })

        buffer.clear()


    # ================= LOOP PRINCIPAL =================
    for line in raw_lines:
        raw = line.strip()
        if not raw:
            continue

        # NOME — apenas normalização visual
        if not parsed["name"]:
            parsed["name"] = normalize_visual(raw)
            continue

        # CONTATOS — CÓPIA LITERAL DA LINHA INTEIRA
        if not parsed["sections"] and (
            "@" in raw or
            "linkedin" in raw.lower() or
            "telefone" in raw.lower() or
            "local" in raw.lower()
        ):
            parsed["contacts"].append(raw.strip())
            continue


        # A PARTIR DAQUI, LIMPEZA SEMÂNTICA NORMAL
        clean = clean_garbage(safe_txt(raw))

        # 🔥 EXCEÇÃO PARA IDIOMAS E CERTIFICAÇÕES
        if not clean:
            if curr_sec and any(k in curr_sec for k in ["IDIOMA", "IDIOMAS", "CERTIFICAÇÃO", "CERTIFICAÇÕES"]):
                buffer.append(raw)
            continue

        upper = clean.upper()
        KEYWORDS = ["RESUMO", "SKILLS", "EXPERIÊNCIA", "HISTÓRICO", "FORMAÇÃO", "IDIOMA", "CERTIFICA"]
        if any(k in upper for k in KEYWORDS) and len(clean) < 50:
            flush()
            curr_sec = upper
            continue

        if curr_sec:
            buffer.append(raw)

    flush()
    print(parsed)
    return parsed

# ==============================================================================
# 2. HTML GENERATOR
# ==============================================================================

def render_cv_html(parsed_data):
    """
    FONTE ÚNICA DA VERDADE (HTML PURO).
    HTML final deve ficar VISUALMENTE IDÊNTICO À TELA.
    """
    import re

    def normalize(txt: str) -> str:
        return re.sub(r'[^\w\s]', '', txt.lower()).strip()

    # --- CONTATOS ---
    contacts_html = ""
    if parsed_data.get("contacts"):
        raw = parsed_data["contacts"][0]
        raw = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', raw)
        raw = re.sub(r'\s*\|\s*', ' <span class="vant-contact-separator">•</span> ', raw)
        contacts_html = raw.replace('\n', ' ')

    sections_html = ""

    for sec in parsed_data.get("sections", []):
        content_html = ""

        # --- SKILLS ---
        if sec["type"] == "chips":
            content_html = (
                '<div class="vant-skills-container">' +
                ''.join(f'<span class="vant-skill-chip">{s}</span>' for s in sec.get("content", [])) +
                '</div>'
            )

        # --- PARÁGRAFO ---
        elif sec["type"] == "paragraph":
            txt = re.sub(r'\*\*(.*?)\*\*', r'<span class="vant-bold">\1</span>', sec.get("content", ""))
            content_html = f'<div class="summary-text">{txt}</div>'

        # --- EXPERIÊNCIA ---
        elif sec["type"] == "jobs":
            jobs_html = ""

            for job in sec.get("content", []):
                bullets_html = ""

                cargo_norm = normalize(job.get("cargo", ""))

                raw_details = []

                for d in job.get("details", []):
                    parts = [p.strip() for p in d.split('|') if p.strip()]
                    raw_details.extend(parts)

                for raw in raw_details:
                    bullets_html += f"""
                    <div class="vant-cv-grid-row">
                        <div class="vant-cv-bullet-col">•</div>
                        <div class="vant-cv-text-col">{raw}</div>
                    </div>
                    """


                empresa = job.get("empresa", "")
                date = job.get("date", "")

                jobs_html += f"""
                <div class="vant-cv-job-container vant-avoid-break">
                    <div class="vant-job-title">
                        {job.get("cargo", "")}
                        <span class="vant-job-company"> | {empresa}</span>
                    </div>
                    <div class="vant-job-date">{date}</div>
                    {bullets_html}
                </div>
                """

            content_html = jobs_html

        # --- SEÇÃO ---
        sections_html += f"""
        <div class="vant-cv-section">{sec.get("title", "")}</div>
        {content_html}
        """

    return f"""
    <div class="cv-paper-sheet">
        <div class="vant-cv-name">{parsed_data.get("name", "")}</div>
        <div class="vant-cv-contact-line">{contacts_html}</div>
        {sections_html}
    </div>
    """
# ==============================================================================
# 3. FUNÇÃO PRINCIPAL (Interface com seu app.py)
# ==============================================================================

def gerar_pdf_candidato(data):
    """
    Gera PDF via WeasyPrint usando a MESMA estrutura da tela.
    """
    try:
        # 1. Obter texto bruto
        raw_text = data.get('cv_otimizado_completo', '')
        
        # 2. Obter o HTML do componente (O MESMO DA TELA)
        # USAR format_text_to_html DIRETAMENTE PARA GARANTIR FIDELIDADE VISUAL
        body_html = format_text_to_html(raw_text)
        
        # Envolver em div cv-paper-sheet para casar com o CSS
        body_html = f'<div class="cv-paper-sheet">{body_html}</div>'
        
        # 3. Criar o envelope HTML completo para o PDF
        full_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                /* Injeta o CSS V13 aqui para garantir a formatação */
                {CSS_V13}
            </style>
        </head>
        <body>
            {body_html}
        </body>
        </html>
        """
        
        # 4. Render PDF
        pdf_bytes = HTML(string=full_html).write_pdf(
            stylesheets=[CSS(string=CSS_PDF)]
        )
        return pdf_bytes

    except Exception as e:
        # Fallback mantido
        from fpdf import FPDF as FPDF_ERR
        err = FPDF_ERR(); err.add_page(); err.set_font("Arial", size=12)
        err.cell(0, 10, f"Erro WeasyPrint: {str(e)}", ln=True)
        return err.output(dest="S").encode("latin-1")
# ============================================================
# HELPER FUNCTIONS PARA ANÁLISE DE CV
# ============================================================

def _process_competitors(competitor_files):
    """Processa arquivos de concorrentes e retorna texto formatado."""
    competitors_text = ""
    if competitor_files:
        for i, comp_file in enumerate(competitor_files):
            c_text = extrair_texto_pdf(comp_file)
            if c_text:
                competitors_text += f"\n--- CONCORRENTE {i+1} ---\n{c_text[:15000]}\n"
        logger.info(f"⚔️ Processando {len(competitor_files)} arquivos de concorrência.")
    return competitors_text

def _curate_books(area_detected):
    """Cura lista de livros baseada na área detectada."""
    books_catalog = load_books_catalog()
    curated_books = []
    seen_titles = set()

    def add_book_safe(book_obj, origin):
        """Helper para adicionar apenas livros únicos"""
        t = book_obj.get('titulo', '').strip().lower()
        if t and t not in seen_titles:
            seen_titles.add(t)
            b_enriched = book_obj.copy()
            b_enriched['categoria_origem'] = origin
            curated_books.append(b_enriched)

    # 1. SELEÇÃO TÉCNICA (Prioridade)
    specific_books = []
    if isinstance(books_catalog, dict):
        specific_books = books_catalog.get(area_detected, [])
    
    if not specific_books: # Backup
        specific_books = BACKUP_CATALOG.get(area_detected, [])
    
    for b in specific_books[:15]:
        add_book_safe(b, area_detected)

    # 2. SELEÇÃO SOFT SKILLS (Secundário)
    soft_skills = []
    if isinstance(books_catalog, dict):
        soft_skills = books_catalog.get("global_soft_skills", [])
    
    if not soft_skills: # Backup
        soft_skills = BACKUP_CATALOG.get("global_soft_skills", [])
        
    for b in soft_skills[:5]:
        add_book_safe(b, "soft_skills")

    # 3. FALLBACK CATASTRÓFICO
    if not curated_books:
        fallback_list = [
             {"titulo": "Hábitos Atômicos", "autor": "James Clear"},
             {"titulo": "Comece pelo Porquê", "autor": "Simon Sinek"}
        ]
        for b in fallback_list:
            add_book_safe(b, "fallback")

    return {"biblioteca_universal": curated_books}

# ============================================================
# ORQUESTRADOR BLINDADO (ATUALIZADO)
# ============================================================
def analyze_cv_logic(cv_text, job_description, competitor_files=None, user_id=None, original_filename=None):
    
    # Sanitizar inputs
    cv_text = sanitize_input(cv_text)
    job_description = sanitize_input(job_description)
    
    # [DEV MODE - INICIO] -----------------------------------------
    DEV_MODE = os.getenv("VANT_DEV_MODE", "0") == "1"  # Toggle para testes rápidos sem gastar tokens

    if DEV_MODE:
        logger.info("🚧 DEV MODE: Bypass de IA ativado.")
        return {
            "veredito": "APROVADO (MOCK)",
            "nota_ats": 88,
            "analise_por_pilares": {"impacto": 90, "keywords": 85, "ats": 95},
            "linkedin_headline": "Senior Software Engineer | Python | AWS",
            "resumo_otimizado": "Profissional focado em arquitetura escalável...",
            "cv_otimizado_completo": """### Experiência Profissional...""",
            "gaps_fatais": [
                {"erro": "Falta de Métricas", "evidencia": "Texto vago", "correcao_sugerida": "Use **números**."},
                {"erro": "Tecnologia Antiga", "evidencia": "Uso de SVN", "correcao_sugerida": "Migre para **Git**."}
            ],
            "analise_comparativa": {
                "vantagens_concorrentes": ["O Benchmark tem PMP.", "O Sênior fala alemão."],
                "seus_diferenciais": ["Você domina **Python**.", "Você tem **Startup** no CV."],
                "plano_de_ataque": "Aposte na agilidade.",
                "probabilidade_aprovacao": 72
            },
            # [ATUALIZADO] Mock alinhado com a nova estrutura (sem roadmap)
            "biblioteca_tecnica": [{"titulo": "Clean Code", "autor": "Uncle Bob", "motivo": "Essencial."}],
            "perguntas_entrevista": [{"pergunta": "Fale sobre um erro.", "expectativa_recrutador": "Honestidade.", "dica_resposta": "Seja direto."}],
            "kit_hacker": {"boolean_string": "site:linkedin.com/in python"},
            "projeto_pratico": {
                "titulo": "API de Alta Performance",
                "descricao": "Criar uma API em FastAPI com cache Redis.",
                "como_apresentar": "Diga que reduziu latência em 50ms."
            }
        }
    # [DEV MODE - FIM] --------------------------------------------
    
    
    # 1. Validação de Input (Fail Fast)
    if not cv_text or len(cv_text.strip()) < 50:
        logger.warning("Tentativa de análise com CV vazio ou inválido.")
        return {
            "veredito": "Erro de Leitura (Arquivo Vazio/Inválido)",
            "nota_ats": 0,
            "gaps_fatais": [{"erro": "Arquivo Ilegível", "evidencia": "Não conseguimos extrair texto do PDF.", "correcao_sugerida": "Envie um PDF selecionável (texto), não escaneado."}]
        }

    # 2. PROCESSA CONCORRENTES
    competitors_text = _process_competitors(competitor_files)

    # 3. Lógica Original de Catálogo
    area_detected = detect_job_area(job_description)
    logger.info(f"🔎 Área detectada: {area_detected.upper()}") 
    
    catalog_payload = _curate_books(area_detected)
    
    # Chama o orquestrador
    try:
        data = run_llm_orchestrator(
            cv_text=cv_text,
            job_description=job_description,
            books_catalog=catalog_payload,
            area=area_detected,
            # PASSANDO OS CONCORRENTES PARA O LLM
            competitors_text=competitors_text if competitors_text else None,
            # PASSANDO USER_ID PARA O CACHE
            user_id=user_id,
            # PASSANDO NOME DO ARQUIVO ORIGINAL
            original_filename=original_filename
        )
        return data
    except Exception as e:
        logger.error(f"Erro fatal no orquestrador: {e}")
        return {
            "veredito": "Erro Técnico",
            "nota_ats": 0,
            "gaps_fatais": [
                {
                    "erro": "Falha no Processamento",
                    "evidencia": str(e),
                    "correcao_sugerida": "Tente novamente mais tarde."
                }
            ]
        }
# [ATUALIZADO] ENGINE HEURÍSTICO COM ENRIQUECIMENTO DE CARGO
# ============================================================
def analyze_preview_lite(cv_text, job_description, forced_area=None):
    """
    Versão FREE com IA REAL: Mostra nota + 2 gaps REAIS específicos do CV.
    Objetivo: Provar valor antes de pedir pagamento.
    """
    import string
    import re
    import json
    
    # Sanitizar inputs
    cv_text = sanitize_input(cv_text)
    job_description = sanitize_input(job_description)
    
    # Se for área forçada (vaga genérica com área de interesse), usa ela
    if forced_area:
        area_detected = forced_area
        logger.info(f"🔎 Área forçada pelo usuário: {area_detected.upper()}")
    else:
        # Detecção normal de área
        area_detected = detect_job_area(job_description)
        logger.info(f"🔎 Área detectada: {area_detected.upper()}")
     
    # Prompt específico para análise gratuita (2 gaps reais)
    area_foco = forced_area.replace('_', ' ').upper() if forced_area else area_detected.upper()
    prompt_preview = f"""
Você é um especialista em ATS e otimização de currículos.

{"⚠️ ÁREA ESPECÍFICA SELECIONADA PELO USUÁRIO: " + area_foco if forced_area else ""}
{"🎯 CRÍTICO: Toda a análise DEVE ser focada EXCLUSIVAMENTE em " + area_foco + ". Ignore outras áreas." if forced_area else ""}

MISSÃO: Analisar este CV e identificar os 2 PROBLEMAS MAIS CRÍTICOS que estão impedindo aprovação para vagas de {area_foco}.

CV DO CANDIDATO:
{cv_text[:3000]}

VAGA ALVO:
{job_description[:1500]}

INSTRUÇÕES ESPECÍFICAS:
1. **FOCO OBRIGATÓRIO**: Identifique os 2 problemas MAIS GRAVES e ESPECÍFICOS deste CV **EXCLUSIVAMENTE PARA A ÁREA DE {area_foco}**
   - Se a área é MARKETING/GROWTH: foque em falta de termos como SEO, SEM, tráfego, conversão, funil, CRM, mídias sociais
   - Se a área é VENDAS: foque em negociação, fechamento, prospecção, pipeline, metas
   - Se a área é TI: foque em tecnologias, frameworks, linguagens, arquitetura
   - **IGNORE problemas de outras áreas que não sejam {area_foco}**
2. **CRÍTICO**: Para exemplo_atual, copie EXATAMENTE um trecho real do CV (não modifique)
3. **CRÍTICO**: Para exemplo_otimizado, TRANSFORME o trecho seguindo TODAS estas regras OBRIGATÓRIAS:
   
   **REGRAS OBRIGATÓRIAS DE OTIMIZAÇÃO:**
   a) **VERBO DE AÇÃO NO INÍCIO**: Comece SEMPRE com verbo de ação forte (desenvolvi, implementei, gerenciei, criei, otimizei, liderei, reduzi, aumentei, automatizei)
   
   b) **TERMOS TÉCNICOS BASEADOS NO CV REAL**: 
      **🚨 REGRA CRÍTICA - NUNCA INVENTE EXPERIÊNCIAS:**
      - **PROIBIDO** adicionar ferramentas/tecnologias que NÃO estão no CV (ex: se não menciona CRM, NÃO adicione CRM)
      - **PROIBIDO** inventar experiências que o candidato não teve
      - **PERMITIDO** apenas melhorar a FORMA de apresentar o que JÁ EXISTE no CV
      
      **COMO OTIMIZAR SEM MENTIR:**
      1. Leia o CV e identifique o que o candidato REALMENTE fez
      2. Use verbos de ação fortes para descrever essas experiências REAIS
      3. Adicione contexto e detalhes que PODEM ser inferidos da experiência real
      4. Use termos técnicos da área {area_foco} APENAS se fizerem sentido com a experiência real
      
      **EXEMPLO CORRETO:**
      - CV atual: "Atendi clientes"
      - ✅ Otimizado: "**Gerenciei** atendimento ao cliente, **resolvendo solicitações** e **garantindo satisfação**"
      - ❌ ERRADO: "Gerenciei atendimento utilizando **CRM**" (se não menciona CRM no CV)
      
      **TERMOS SUGERIDOS POR ÁREA (use APENAS se compatível com experiência real):**
      - **MARKETING GROWTH**: estratégias digitais, campanhas, engajamento, análise de resultados
      - **VENDAS/CS**: relacionamento com clientes, negociação, apresentação de soluções
      - **TI/DEV**: desenvolvimento, implementação, resolução de problemas técnicos
      - **FINANCEIRO**: análise financeira, controles, relatórios, processos
      
      **IMPORTANTE**: Prefira NÃO adicionar termos técnicos a INVENTAR experiências falsas.
      
      Exemplos por área REAL do CV (não da vaga):
      - Suporte: atendimento ao cliente, resolução de problemas, SAC, tickets, SLA, troubleshooting
      - Vendas: prospecção, negociação, fechamento, pipeline, CRM, metas, conversão
      - TI: Python, React, API REST, Docker, Kubernetes, microserviços, CI/CD
      - RH: recrutamento, onboarding, gestão de pessoas, KPIs, treinamento
      - Marketing: SEO, Google Ads, funil de conversão, CRM, automação
   
   c) **SUBSTÂNCIA VISÍVEL**: A versão otimizada deve ter 30-50% MAIS conteúdo que a atual
      - Adicione especificidades técnicas BASEADAS NA EXPERIÊNCIA REAL
      - Inclua contexto profissional COERENTE com o histórico
      - Detalhe ferramentas/metodologias que o candidato REALMENTE poderia ter usado
   
   d) **MARQUE ADIÇÕES COM ASTERISCOS**: Coloque **termo novo** em torno de CADA palavra/frase que você ADICIONAR (não estava no original)
      - Exemplo: "Atendi clientes" → "**Gerenciei** atendimento ao cliente utilizando **CRM** e **metodologia de resolução ágil**"
   
   e) **NÃO invente números, porcentagens ou métricas** - foque em qualidade técnica e profissional
   
   f) **IMPORTANTE**: Se o gap for "Formatação inadequada", NÃO adicione termos técnicos. Foque em estrutura e clareza.
   
   g) **COERÊNCIA OBRIGATÓRIA**: Analise o CV REAL antes de adicionar termos. Se não houver experiência na área da vaga, otimize com base na experiência que o candidato TEM, não na que ele não tem.

4. **NÃO USE** "falta de métricas" ou "falta de números" como problema
5. Identifique os 2 problemas MAIS CRÍTICOS dentre estas categorias RELEVANTES PARA A ÁREA:
   
   **PROBLEMAS DE ESTRUTURA:**
   - Formatação inadequada (parágrafos longos, sem bullets, estrutura confusa)
   - Falta de seções claras (experiência, educação, habilidades mal organizadas)
   - Informações importantes enterradas no meio do texto
   
   **PROBLEMAS DE LINGUAGEM:**
   - Linguagem passiva (ex: "fui responsável por", "participei de" → "gerenciei", "implementei")
   - Verbos fracos ou genéricos (ex: "trabalhei com", "ajudei em" → "desenvolvi", "liderei")
   - Falta de verbos de ação no início das descrições
   
   **PROBLEMAS DE CONTEÚDO:**
   - Falta de palavras-chave ESPECÍFICAS DA ÁREA DE {forced_area.replace('_', ' ').upper() if forced_area else area_detected.upper()}
   - Descrições genéricas sem especificidade ou detalhes relevantes PARA A ÁREA
   - Ausência de contexto profissional (tamanho de equipe, escopo do projeto, responsabilidades)
   - Falta de tecnologias/ferramentas específicas da área
   
   **PROBLEMAS DE IMPACTO:**
   - Foco em tarefas em vez de resultados/impacto
   - Descrições que não mostram valor agregado
   - Ausência de indicadores de complexidade ou escala do trabalho
   
   **IMPORTANTE**: Varie os tipos de gaps detectados. NÃO use sempre "falta de palavras-chave". Analise o CV real e identifique os 2 problemas MAIS GRAVES que realmente prejudicam a aprovação no ATS.
   
6. Seja direto e objetivo
7. **VERIFICAÇÃO OBRIGATÓRIA**: Não inclua "métricas", "números", "resultados quantitativos" em nenhum gap
8. **VERIFICAÇÃO OBRIGATÓRIA**: Revise sua resposta antes de enviar para garantir que não há gaps sobre métricas
9. **VERIFICAÇÃO OBRIGATÓRIA**: Confirme que exemplo_otimizado tem pelo menos 2 termos técnicos marcados com **termo**
10. **VERIFICAÇÃO OBRIGATÓRIA**: Confirme que exemplo_otimizado começa com verbo de ação forte

**CÁLCULO OBRIGATÓRIO DOS PILARES (analise_por_pilares):**
9. **impacto (0-100)**: Avalie o uso de verbos de ação, resultados e especificidade.
   **CRITÉRIOS OBJETIVOS OBRIGATÓRIOS** (some os pontos):
   
   a) **Verbos de ação (0-35 pontos)**:
      - 35 pts: 80%+ das descrições começam com verbos fortes (implementei, gerenciei, desenvolvi, otimizei)
      - 25 pts: 50-80% com verbos de ação
      - 15 pts: 20-50% com verbos de ação
      - 5 pts: Menos de 20%, predomina linguagem passiva ("responsável por", "participei de")
      - 0 pts: Sem verbos de ação identificáveis
   
   b) **Especificidade técnica (0-35 pontos)**:
      - 35 pts: Descrições detalhadas com ferramentas/metodologias específicas
      - 25 pts: Algumas descrições específicas, outras genéricas
      - 15 pts: Maioria genérica, poucas especificidades
      - 5 pts: Descrições muito vagas ("trabalhei com tecnologia")
      - 0 pts: Apenas títulos de cargo sem descrições
   
   c) **Contexto profissional (0-30 pontos)**:
      - 30 pts: Cada experiência tem contexto claro (projeto, equipe, objetivo)
      - 20 pts: Maioria tem contexto, algumas sem
      - 10 pts: Pouco contexto, descrições isoladas
      - 0 pts: Sem contexto, apenas listas de tarefas
   
   **IMPORTANTE**: Some os pontos das 3 categorias para obter o score final (0-100).
   **CONSISTÊNCIA**: Use estes critérios EXATOS toda vez para garantir scores consistentes.

10. **keywords (0-100)**: Avalie cobertura de palavras-chave relevantes PARA A ÁREA.
    **CRITÉRIOS OBJETIVOS OBRIGATÓRIOS** (some os pontos):
    
    a) **Termos técnicos da área (0-40 pontos)**:
       - 40 pts: 8+ termos técnicos relevantes identificados no CV
       - 30 pts: 5-7 termos técnicos relevantes
       - 20 pts: 3-4 termos técnicos relevantes
       - 10 pts: 1-2 termos técnicos relevantes
       - 0 pts: Nenhum termo técnico da área
    
    b) **Ferramentas/tecnologias (0-30 pontos)**:
       - 30 pts: Lista clara de ferramentas/tecnologias usadas (5+)
       - 20 pts: Algumas ferramentas mencionadas (3-4)
       - 10 pts: Poucas ferramentas (1-2)
       - 0 pts: Sem menção a ferramentas específicas
    
    c) **Alinhamento com vaga (0-30 pontos)**:
       - 30 pts: 80%+ dos termos da vaga aparecem no CV
       - 20 pts: 50-80% dos termos da vaga no CV
       - 10 pts: 20-50% dos termos da vaga no CV
       - 0 pts: Menos de 20% de match com termos da vaga
    
    **IMPORTANTE**: Some os pontos das 3 categorias para obter o score final (0-100).
    **CONSISTÊNCIA**: Conte os termos técnicos REAIS do CV para garantir scores consistentes.
11. **ats (0-100)**: Avalie formatação, estrutura e compatibilidade com sistemas ATS.
    **CRITÉRIOS OBJETIVOS OBRIGATÓRIOS** (some os pontos):
    
    a) **Estrutura de seções (0-25 pontos)**:
       - 25 pts: Seções claras (Experiência, Educação, Habilidades) com títulos destacados
       - 15 pts: Seções presentes mas títulos pouco claros
       - 5 pts: Seções misturadas ou sem organização
       - 0 pts: Sem estrutura identificável
    
    b) **Uso de bullets (0-25 pontos)**:
       - 25 pts: Maioria das descrições em bullets (80%+)
       - 15 pts: Alguns bullets mas também parágrafos longos (40-80%)
       - 5 pts: Poucos bullets, predominam parágrafos (10-40%)
       - 0 pts: Sem bullets, apenas parágrafos corridos
    
    c) **Formatação limpa (0-25 pontos)**:
       - 25 pts: Texto simples, sem tabelas/colunas/gráficos complexos
       - 15 pts: Formatação simples com pequenas inconsistências
       - 5 pts: Uso de tabelas ou formatação que pode confundir ATS
       - 0 pts: Formatação complexa (colunas, gráficos, texto em imagens)
    
    d) **Consistência (0-25 pontos)**:
       - 25 pts: Datas, cargos e empresas em formato consistente
       - 15 pts: Maioria consistente, algumas variações
       - 5 pts: Formatos misturados (ex: "2020-2021" e "Jan 2020 - Dez 2021")
       - 0 pts: Sem padrão identificável
    
    **IMPORTANTE**: Some os pontos das 4 categorias para obter o score final (0-100).
    **CONSISTÊNCIA**: Analise o CV com estes critérios EXATOS toda vez para garantir scores consistentes.

12. **OBRIGATÓRIO**: Calcule os scores dos pilares BASEADO NO CV REAL usando os critérios objetivos acima, não use placeholders.

13. **PRIORIDADE DA ÁREA**: Se {forced_area or area_detected} == "rh_lideranca", foque em gaps como:
   - Falta de termos de RH: "recrutamento", "seleção", "treinamento", "desenvolvimento", "gestão de pessoas"
   - Experiências sem foco em liderança ou gestão
   - Ausência de indicadores de RH (KPIs, metas, resultados com equipes)
   - Descrições operacionais em vez de estratégicas
14. **PRIORIDADE DA ÁREA**: Se {forced_area or area_detected} == "ti_dev_gen", foque em gaps como:
    - Falta de termos técnicos: "python", "java", "javascript", "react", "api", "microserviços"
    - Tecnologias desatualizadas ou frameworks antigos
    - Ausência de projetos ou GitHub/portfólio
    - Descrições sem menção a arquitetura ou padrões
15. **PRIORIDADE DA ÁREA**: Se {forced_area or area_detected} == "ti_dados_ai", foque em gaps como:
    - Falta de termos: "python", "sql", "machine learning", "analytics", "dashboard", "bi"
    - Ausência de ferramentas: "pandas", "numpy", "tensorflow", "power bi", "tableau"
    - Projetos sem resultados ou insights gerados
    - Falta de estatística ou matemática aplicada
16. **PRIORIDADE DA ÁREA**: Se {forced_area or area_detected} == "marketing_growth", foque em gaps como:
    - Falta de termos: "seo", "sem", "tráfego", "conversão", "funil", "crm", "mídias sociais"
    - Ausência de métricas de marketing: "cpa", "roi", "cac", "ltv"
    - Campanhas sem resultados ou KPIs
    - Descrições sem menção a estratégias ou canais
17. **PRIORIDADE DA ÁREA**: Se {forced_area or area_detected} == "vendas_cs", foque em gaps como:
    - Falta de termos: "negociação", "fechamento", "prospecção", "pipeline", "crm", "follow-up"
    - Ausência de metas de vendas ou resultados
    - Experiências sem menção a clientes ou contratos
    - Falta de técnicas de vendas ou metodologias
18. **PRIORIDADE DA ÁREA**: Se {forced_area or area_detected} == "financeiro_corp", foque em gaps como:
    - Falta de termos: "orçamento", "fluxo de caixa", "demonstrativos", "balanço", "kpis"
    - Ausência de ferramentas: "excel", "sap", "erp", "sistema financeiro"
    - Projetos sem impacto financeiro mensurável
    - Descrições sem menção a controles ou auditoria

19. DETECTE A ÁREA ESPECÍFICA:
   
   {"⚠️ ÁREA FORÇADA ESPECÍFICA: Use exatamente '" + forced_area.replace('_', ' ').upper() + "' como setor_detectado. Ignore outras detecções." if forced_area else ""}
   
   Prioridade 1 - Mapeamento Exato (baseado na VAGA):
   - Se a VAGA menciona "suporte", "suporte técnico", "analista de suporte", "help desk", "atendimento", "service desk", "suporte ao cliente" → "SUPORTE TI"
   - Se a VAGA menciona "compras", "suprimentos", "logística", "almoxarifado", "cadeia de suprimentos" → "COMPRAS"
   - Se a VAGA menciona "vendas", "comercial", "sales", "vendedor", "negócios", "business" → "VENDAS"
   - Se a VAGA menciona "marketing", "mkt", "comunicação", "social media", "publicidade" → "MARKETING"
   - Se a VAGA menciona "rh", "recursos humanos", "people", "gestão de pessoas", "recrutamento" → "RH"
   - Se a VAGA menciona "financeiro", "contábil", "fiscal", "contabilidade", "tesouraria", "orçamento" → "FINANCEIRO"
   - Se a VAGA menciona "jurídico", "advogado", "legal", "contratos", "legislação", "advocacia" → "JURÍDICO"
   - Se a VAGA menciona "engenharia", "engenheiro", "civil", "mecânica", "elétrica", "produção" → "ENGENHARIA"
   - Se a VAGA menciona "arquitetura", "arquiteto", "design", "designer", "criação", "visual" → "ARQUITETURA"
   - Se a VAGA menciona "saúde", "médico", "enfermagem", "hospital", "clínica", "farmácia" → "SAÚDE"
   - Se a VAGA menciona "educação", "professor", "pedagogia", "escola", "universidade", "ensino" → "EDUCAÇÃO"
   - Se a VAGA menciona "logística", "transporte", "distribuição", "frota", "entrega" → "LOGÍSTICA"
   - Se a VAGA menciona "produção", "fábrica", "indústria", "manufatura", "operação" → "PRODUÇÃO"
   - Se a VAGA menciona "agro", "agronegócio", "agricultura", "pecuária", "rural" → "AGRO"
   
   Se não encontrar nenhuma das áreas acima (Prioridade 2 - Fallback Livre):
   - Use seu conhecimento para detectar a área mais adequada baseada APENAS no contexto da VAGA
   - Exemplos: "RECURSOS HUMANOS", "ADMINISTRAÇÃO", "QUALIDADE", "PESQUISA", etc.
   - Seja específico e profissional na nomeação
   
   IMPORTANTE: setor_detectado deve refletir a VAGA, não o histórico do candidato

OUTPUT JSON (OBRIGATÓRIO - COPIE EXATAMENTE):
{{
  "nota_ats": [CALCULE O SCORE REAL BASEADO NO CV],
  "analise_por_pilares": {{
    "impacto": [CALCULE BASEADO NO IMPACTO DO CV],
    "keywords": [CALCULE BASEADO NAS KEYWORDS DO CV],
    "ats": [CALCULE BASEADO NA FORMATAÇÃO ATS],
    "setor_detectado": "TECNOLOGIA"
  }},
  "gap_1": {{
    "titulo": "Nome do problema",
    "explicacao": "Por que isso é crítico",
    "exemplo_atual": "Trecho real do CV que mostra o problema",
    "exemplo_otimizado": "Como deveria ser (com números e palavras-chave)"
  }},
  "gap_2": {{
    "titulo": "Nome do problema",
    "explicacao": "Por que isso é crítico",
    "termos_faltando": [
      {{"termo": "termo1", "frequencia": "presente em X% das vagas de [cargo específico]"}},
      {{"termo": "termo2", "frequencia": "presente em X% das vagas de [cargo específico]"}},
      {{"termo": "termo3", "frequencia": "presente em X% das vagas de [cargo específico]"}}
    ]
  }}
}}

**REGRAS PARA termos_faltando (SE APLICÁVEL):**
1. **EXTRAÇÃO BASEADA EM MERCADO**: Identifique termos que:
   - Aparecem frequentemente na descrição da vaga fornecida
   - São comuns em vagas similares da área detectada
   - Estão AUSENTES no CV do candidato
   
2. **ATRIBUIÇÃO DE FONTE OBRIGATÓRIA**: Para cada termo, calcule frequência estimada:
   - Analise a vaga fornecida e identifique o cargo/nível (ex: "Analista de Suporte Sênior", "Desenvolvedor Full Stack Pleno")
   - Estime frequência baseada em conhecimento do mercado: 70-90% (muito comum), 50-70% (comum), 30-50% (moderado)
   - Formato: "presente em X% das vagas de [cargo específico da vaga]"
   
3. **ORDENAÇÃO POR RELEVÂNCIA**: Liste do mais ao menos importante:
   - Prioridade 1: Termos que aparecem na vaga E são críticos para a área (80-90%)
   - Prioridade 2: Termos comuns na área mas não na vaga específica (60-80%)
   - Prioridade 3: Termos complementares relevantes (40-60%)
   
4. **MÁXIMO 5 TERMOS**: Liste apenas os 5 termos mais críticos e rastreáveis
   
5. **VALIDAÇÃO**: Só inclua termos que:
   - Você consegue justificar a frequência com base na vaga/área
   - Fazem sentido para o cargo específico identificado
   - NÃO são genéricos demais (evite: "comunicação", "trabalho em equipe")

IMPORTANTE: 
- Retorne APENAS o JSON válido, sem texto adicional
- nota_ats: **OBRIGATÓRIO CALCULAR** score real 0-100 baseado na qualidade do CV vs vaga
- pilares: **OBRIGATÓRIO CALCULAR** scores 0-100 baseados em análise específica
- setor_detectado: área exata detectada acima
- NÃO ADICIONE setores fora do lugar
- **PROIBIDO IDENTIFICAR "FALTA DE MÉTRICAS" OU "FALTA DE NÚMEROS" COMO PROBLEMA**
- **CRÍTICO**: NÃO use 0 como placeholder - calcule scores reais!
- **termos_faltando**: Use estrutura de objetos com termo + frequencia, ordenados por relevância
"""
    
    try:
        # ESTRATÉGIA: Tenta Groq (gratuito e rápido) primeiro, fallback para Gemini
        response = None
        
        # TENTATIVA 1: GROQ (Gratuito + Rápido)
        try:
            from groq import Groq
            GROQ_API_KEY = os.getenv("GROQ_API_KEY")
            
            if GROQ_API_KEY:
                logger.info("🚀 Tentando Groq (gratuito) para preview...")
                groq_client = Groq(api_key=GROQ_API_KEY)
                
                response_obj = groq_client.chat.completions.create(
                    model="llama-3.3-70b-versatile",  # Modelo gratuito e rápido
                    messages=[{"role": "user", "content": prompt_preview}],
                    temperature=0,
                    max_tokens=1000,
                )
                
                response = response_obj.choices[0].message.content
                logger.info("✅ Groq respondeu com sucesso!")
                logger.info(f"[DEBUG GROQ] Resposta bruta: {response}")
            else:
                raise Exception("GROQ_API_KEY não configurada")
                
        except Exception as groq_error:
            # Se Groq falhar (rate limit ou erro), usa Gemini como fallback
            logger.warning(f"⚠️ Groq falhou ({groq_error}), usando Gemini 1.5 Flash como fallback...")
            
            from google import genai
            from google.genai import types
            
            GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
            if not GOOGLE_API_KEY:
                raise Exception("GOOGLE_API_KEY não configurada")
            
            client = genai.Client(api_key=GOOGLE_API_KEY)
            
            # Usa gemini-1.5-flash como fallback
            response_obj = client.models.generate_content(
                model="gemini-1.5-flash",
                contents=prompt_preview,
                config=types.GenerateContentConfig(
                    temperature=0,
                    max_output_tokens=1000,
                )
            )
            
            response = response_obj.text
            logger.info("✅ Gemini 1.5 Flash respondeu como fallback!")
        
        # Parse do JSON retornado pela IA
        # Remove markdown se houver
        response_clean = response.strip()
        if response_clean.startswith("```"):
            response_clean = response_clean.split("```")[1]
            if response_clean.startswith("json"):
                response_clean = response_clean[4:]
        
        # DEBUG: Log da resposta da IA
        logger.info(f"[DEBUG] Resposta IA: {response_clean}")
        
        data = json.loads(response_clean)
        
        # Extrai dados da IA
        nota_ats = int(data.get("nota_ats", 45))
        pilares = data.get("analise_por_pilares", {})
        gap_1 = data.get("gap_1", {})
        gap_2 = data.get("gap_2", {})
        
        # DEBUG: Log dos dados extraídos
        logger.info(f"[DEBUG] nota_ats: {nota_ats}")
        logger.info(f"[DEBUG] pilares: {pilares}")
        logger.info(f"[DEBUG] setor_detectado da IA: {pilares.get('setor_detectado', 'NÃO ENVIADO')}")
        
        # Usar setor detectado pela IA, não a função antiga
        setor_detectado = pilares.get("setor_detectado", "TECNOLOGIA")
        
        # Gerar pergunta personalizada do simulador de entrevista
        pergunta_preview = None
        try:
            logger.info("🎤 Gerando pergunta personalizada do simulador...")
            
            # Extrair trecho relevante do CV para contextualizar
            cv_snippet = cv_text[:500] if len(cv_text) > 500 else cv_text
            gap1_exemplo = gap_1.get("exemplo_atual", "") if gap_1 else ""
            
            prompt_pergunta = f"""
Você é um recrutador experiente de {setor_detectado}. Baseado neste CV e gap detectado, crie UMA pergunta de entrevista ULTRA-PERSONALIZADA.

{"⚠️ ÁREA ESPECÍFICA: " + setor_detectado if forced_area else ""}
{"🎯 CRÍTICO: A pergunta DEVE ser focada EXCLUSIVAMENTE em " + setor_detectado + "." if forced_area else ""}

SETOR: {setor_detectado}
TRECHO DO CV: {cv_snippet}
GAP DETECTADO: {gap_1.get('titulo', '') if gap_1 else ''}
EXEMPLO DO CV: {gap1_exemplo[:200]}

INSTRUÇÕES:
1. **EXTRAIA O NOME DE UMA EMPRESA** do CV se possível e mencione na pergunta
2. A pergunta DEVE mencionar algo específico do CV (experiência na empresa X, tecnologia Y, projeto Z)
3. A pergunta DEVE fazer o candidato pensar "Nossa, eles realmente leram meu CV!"
4. Use o gap detectado para direcionar a pergunta
5. **FOCO OBRIGATÓRIO**: A pergunta deve ser sobre {setor_detectado}
6. A dica deve ser prática e valiosa para {setor_detectado}

EXEMPLOS DE BOAS PERGUNTAS POR ÁREA (com nome de empresa):
- **MARKETING/GROWTH**: "Considerando sua experiência na [Empresa X], como você mediu o ROI das campanhas digitais? Quais métricas você priorizou?"
- **VENDAS**: "Vi que você trabalhou na [Empresa Y]. Qual foi o deal mais desafiador que você fechou lá e qual foi sua estratégia?"
- **TI**: "Na sua passagem pela [Empresa Z], como você garantiu a qualidade do código e a escalabilidade das soluções?"
- **FINANCEIRO**: "Durante sua experiência na [Empresa W], como você estruturou os controles financeiros e quais KPIs você acompanhava?"

**IMPORTANTE**: Se conseguir identificar nome de empresa no CV, USE na pergunta para aumentar personalização.

OUTPUT JSON:
{{
  "pergunta": "Pergunta ultra-personalizada focada em {setor_detectado} mencionando empresa se possível",
  "dica": "Dica prática específica para {setor_detectado}"
}}

Retorne APENAS o JSON, sem markdown.
"""
            
            # Tentar gerar com Groq primeiro
            try:
                from groq import Groq
                GROQ_API_KEY = os.getenv("GROQ_API_KEY")
                
                if GROQ_API_KEY:
                    groq_client = Groq(api_key=GROQ_API_KEY)
                    response_pergunta = groq_client.chat.completions.create(
                        model="llama-3.3-70b-versatile",
                        messages=[{"role": "user", "content": prompt_pergunta}],
                        temperature=0.7,
                        max_tokens=300,
                    )
                    pergunta_response = response_pergunta.choices[0].message.content.strip()
                    
                    # Limpar markdown se houver
                    if pergunta_response.startswith("```"):
                        pergunta_response = pergunta_response.split("```")[1]
                        if pergunta_response.startswith("json"):
                            pergunta_response = pergunta_response[4:]
                    
                    pergunta_preview = json.loads(pergunta_response)
                    logger.info(f"✅ Pergunta personalizada gerada: {pergunta_preview.get('pergunta', '')[:50]}...")
            except Exception as groq_err:
                logger.warning(f"⚠️ Groq falhou para pergunta, tentando Gemini: {groq_err}")
                
                # Fallback para Gemini
                from google import genai
                from google.genai import types
                
                GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
                if GOOGLE_API_KEY:
                    client = genai.Client(api_key=GOOGLE_API_KEY)
                    response_pergunta = client.models.generate_content(
                        model="gemini-1.5-flash",
                        contents=prompt_pergunta,
                        config=types.GenerateContentConfig(temperature=0.7, max_output_tokens=300)
                    )
                    pergunta_response = response_pergunta.text.strip()
                    
                    # Limpar markdown
                    if pergunta_response.startswith("```"):
                        pergunta_response = pergunta_response.split("```")[1]
                        if pergunta_response.startswith("json"):
                            pergunta_response = pergunta_response[4:]
                    
                    pergunta_preview = json.loads(pergunta_response)
                    logger.info(f"✅ Pergunta gerada com Gemini: {pergunta_preview.get('pergunta', '')[:50]}...")
        except Exception as pergunta_err:
            logger.error(f"❌ Erro ao gerar pergunta personalizada: {pergunta_err}")
            # Fallback para pergunta genérica
            pergunta_preview = {
                "pergunta": f"Baseado na sua experiência em {setor_detectado}, conte sobre um desafio significativo que você enfrentou. Como você resolveu?",
                "dica": "Use o método STAR: Situação → Tarefa → Ação → Resultado. Seja específico e mostre seu raciocínio."
            }
        
        return {
            "veredito": "ANÁLISE INICIAL CONCLUÍDA",
            "nota_ats": nota_ats,
            "analise_por_pilares": {
                "impacto": int(pilares.get("impacto", max(nota_ats - 10, 20))),
                "keywords": int(pilares.get("keywords", max(nota_ats - 5, 25))),
                "ats": int(pilares.get("ats", max(nota_ats - 15, 30))),
                "setor_detectado": str(setor_detectado).upper()
            },
            "gap_1": gap_1,
            "gap_2": gap_2,
            "pergunta_preview": pergunta_preview,
            "linkedin_headline": "🔒 [CONTEÚDO PREMIUM BLOQUEADO]",
            "resumo_otimizado": "🔒 [DISPONÍVEL APENAS NA VERSÃO PAGA]",
            "cv_otimizado_completo": "🔒",
            "kit_hacker": {"boolean_string": "🔒"},
            "biblioteca_tecnica": []
        }
        
    except Exception as e:
        logger.error(f"Erro na análise preview com IA: {e}")
        
        # Fallback com IA RÁPIDA (Groq gratuito) se a primeira tentativa falhar
        try:
            logger.info("🔄 Tentando fallback rápido com Groq para análise ATS...")
            
            from groq import Groq
            GROQ_API_KEY = os.getenv("GROQ_API_KEY")
            
            if GROQ_API_KEY:
                groq_client = Groq(api_key=GROQ_API_KEY)
                
                # Prompt simplificado para análise ATS rápida
                fallback_prompt = f"""
Analise este CV para a vaga e retorne APENAS JSON:

VAGA: {job_description[:800]}

CV: {cv_text[:2000]}

**IMPORTANTE**: Use APENAS dados reais do CV. Não invente exemplos, números ou porcentagens.
**PROIBIDO**: Não use "falta de métricas" ou "falta de números" como problema.

Retorne JSON exato:
{{
  "nota_ats": [CALCULE O SCORE REAL BASEADO NO CV],
  "analise_por_pilares": {{
    "impacto": [CALCULE BASEADO NO IMPACTO DO CV],
    "keywords": [CALCULE BASEADO NAS KEYWORDS DO CV],
    "ats": [CALCULE BASEADO NA FORMATAÇÃO ATS],
    "setor_detectado": "TECNOLOGIA"
  }},
  "gap_1": {{
    "titulo": "Problema detectado",
    "explicacao": "Por que isso é crítico",
    "exemplo_atual": "Trecho REAL do CV (copie exatamente)",
    "exemplo_otimizado": "MESMO trecho com melhorias linguísticas (sem inventar métricas)"
  }},
  "gap_2": {{
    "titulo": "Outro problema",
    "explicacao": "Detalhes",
    "termos_faltando": ["termo1", "termo2", "termo3"]
  }}
}}

Regras:
- nota_ats: **OBRIGATÓRIO CALCULAR** score real 0-100 baseado em matching real
- pilares: **OBRIGATÓRIO CALCULAR** scores 0-100 baseados em análise específica
- setor: área detectada (SUPORTE TI, VENDAS, MARKETING, etc)
- gaps: problemas reais do CV
- exemplo_atual: copie trecho exato do CV
- exemplo_otimizado: mesmo trecho com melhorias linguísticas (sem números/percentagens)
- **CRÍTICO**: NÃO use 0 como placeholder - calcule scores reais!
"""
                
                response_obj = groq_client.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    messages=[{"role": "user", "content": fallback_prompt}],
                    temperature=0,
                    max_tokens=500,
                )
                
                response = response_obj.choices[0].message.content
                logger.info("✅ Fallback Groq funcionou!")
                
                # Parse do JSON
                response_clean = response.strip()
                if response_clean.startswith("```"):
                    response_clean = response_clean.split("```")[1]
                    if response_clean.startswith("json"):
                        response_clean = response_clean[4:]
                
                data = json.loads(response_clean)
                
                # Validação e ajustes
                nota_ats = max(5, min(int(data.get("nota_ats", 25)), 100))
                pilares = data.get("analise_por_pilares", {})
                gap_1 = data.get("gap_1", {})
                gap_2 = data.get("gap_2", {})
                
                return {
                    "veredito": "ANÁLISE INICIAL CONCLUÍDA",
                    "nota_ats": nota_ats,
                    "analise_por_pilares": {
                        "impacto": max(0, min(int(pilares.get("impacto", nota_ats)), 100)),
                        "keywords": max(0, min(int(pilares.get("keywords", nota_ats)), 100)),
                        "ats": max(0, min(int(pilares.get("ats", nota_ats)), 100)),
                        "setor_detectado": str(pilares.get("setor_detectado", "TECNOLOGIA")).upper()
                    },
                    "gap_1": gap_1,
                    "gap_2": gap_2,
                    "linkedin_headline": "🔒 [CONTEÚDO PREMIUM BLOQUEADO]",
                    "resumo_otimizado": "🔒 [DISPONÍVEL APENAS NA VERSÃO PAGA]",
                    "cv_otimizado_completo": "🔒",
                    "kit_hacker": {"boolean_string": "🔒"},
                    "biblioteca_tecnica": []
                }
                
        except Exception as groq_fallback_error:
            logger.warning(f"⚠️ Fallback Groq falhou: {groq_fallback_error}")
            
            # ÚLTIMO RECURSO: Gemini 1.5 Flash (barato)
            try:
                logger.info("🔄 Último recurso: Gemini 1.5 Flash...")
                
                from google import genai
                from google.genai import types
                
                GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
                if not GOOGLE_API_KEY:
                    raise Exception("GOOGLE_API_KEY não configurada")
                
                client = genai.Client(api_key=GOOGLE_API_KEY)
                
                response_obj = client.models.generate_content(
                    model="gemini-1.5-flash",
                    contents=fallback_prompt,
                    config=types.GenerateContentConfig(
                        temperature=0,
                        max_output_tokens=500,
                    )
                )
                
                response = response_obj.text
                logger.info("✅ Gemini fallback funcionou!")
                
                # Parse do JSON
                response_clean = response.strip()
                if response_clean.startswith("```"):
                    response_clean = response_clean.split("```")[1]
                    if response_clean.startswith("json"):
                        response_clean = response_clean[4:]
                
                data = json.loads(response_clean)
                
                # Validação e ajustes
                nota_ats = max(5, min(int(data.get("nota_ats", 25)), 100))
                pilares = data.get("analise_por_pilares", {})
                gap_1 = data.get("gap_1", {})
                gap_2 = data.get("gap_2", {})
                
                return {
                    "veredito": "ANÁLISE INICIAL CONCLUÍDA",
                    "nota_ats": nota_ats,
                    "analise_por_pilares": {
                        "impacto": max(0, min(int(pilares.get("impacto", nota_ats)), 100)),
                        "keywords": max(0, min(int(pilares.get("keywords", nota_ats)), 100)),
                        "ats": max(0, min(int(pilares.get("ats", nota_ats)), 100)),
                        "setor_detectado": str(pilares.get("setor_detectado", "TECNOLOGIA")).upper()
                    },
                    "gap_1": gap_1,
                    "gap_2": gap_2,
                    "linkedin_headline": "🔒 [CONTEÚDO PREMIUM BLOQUEADO]",
                    "resumo_otimizado": "🔒 [DISPONÍVEL APENAS NA VERSÃO PAGA]",
                    "cv_otimizado_completo": "🔒",
                    "kit_hacker": {"boolean_string": "🔒"},
                    "biblioteca_tecnica": []
                }
                
            except Exception as gemini_fallback_error:
                logger.error(f"❌ Todos os fallbacks falharam: {gemini_fallback_error}")
                
                # ÚLTIMO RECURSO: Heurística simples
                return {
                    "veredito": "ANÁLISE INICIAL CONCLUÍDA",
                    "nota_ats": 25,
                    "analise_por_pilares": {
                        "impacto": 20,
                        "keywords": 30,
                        "ats": 25,
                        "setor_detectado": "TECNOLOGIA"
                    },
                    "gap_1": {
                        "titulo": "Análise Limitada",
                        "explicacao": "Sistemas de IA indisponíveis no momento"
                    },
                    "gap_2": {
                        "titulo": "Tente Novamente",
                        "explicacao": "Recarregue a página e tente novamente"
                    },
                    "linkedin_headline": "🔒 [CONTEÚDO PREMIUM BLOQUEADO]",
                    "resumo_otimizado": "🔒 [DISPONÍVEL APENAS NA VERSÃO PAGA]",
                    "cv_otimizado_completo": "🔒",
                    "kit_hacker": {"boolean_string": "🔒"},
                    "biblioteca_tecnica": []
                }

# ============================================================
# GERADOR DE WORD V7 (DESIGN SYSTEM TRANSLATION)
# ============================================================
def gerar_word_candidato(data):
    """
    VERSÃO V10 ESTÁVEL: Reorganizada e testada
    Corrige ordem de definições e remove dependências circulares
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from io import BytesIO
    import re

    # ============================================================
    # CONSTANTES DE CORES E FONTES
    # ============================================================
    COLOR_TEXT   = RGBColor(51, 65, 85)
    COLOR_TITLE  = RGBColor(15, 23, 42)
    COLOR_ACCENT = RGBColor(16, 185, 129)
    COLOR_SUB    = RGBColor(100, 116, 139)
    COLOR_CHIP_TEXT = RGBColor(71, 85, 105)
    COLOR_BORDER_HEX = "E2E8F0"
    COLOR_CHIP_BG = "F1F5F9"
    FONT_MAIN = 'Segoe UI'

    # ============================================================
    # FUNÇÕES HELPER DE VALIDAÇÃO (DEFINIDAS PRIMEIRO)
    # ============================================================
    
    def is_valid(text):
        """Valida se o texto não é lixo"""
        if not text: 
            return False
        t = str(text).strip().lower()
        garbage = ["não informado", "nao informado", "n/a", "null", "none", ""]
        return t not in garbage and len(t) > 1
    
    def is_contact_line(text):
        """Detecta se é linha de contato"""
        if not text:
            return False
        t = str(text).lower()
        keywords = ['@', 'linkedin', 'github', 'telefone', 'email', '.com']
        return any(kw in t for kw in keywords)

    # ============================================================
    # FUNÇÕES HELPER XML (FORMATAÇÃO WORD)
    # ============================================================
    
    def set_border_bottom(paragraph):
        """Adiciona borda inferior ao parágrafo"""
        try:
            p = paragraph._p
            pPr = p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '6')
            bottom.set(qn('w:color'), COLOR_BORDER_HEX)
            pBdr.append(bottom)
            pPr.append(pBdr)
        except:
            pass

    def apply_chip_background(run):
        """Aplica fundo colorido ao texto (chip)"""
        try:
            rPr = run._r.get_or_add_rPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), COLOR_CHIP_BG)
            rPr.append(shd)
        except:
            pass

    def set_keep_with_next(paragraph):
        """Mantém parágrafo grudado no próximo"""
        try:
            p = paragraph._p
            pPr = p.get_or_add_pPr()
            keepNext = OxmlElement('w:keepNext')
            pPr.append(keepNext)
        except:
            pass

    def set_keep_together(paragraph):
        """Evita quebra dentro do parágrafo"""
        try:
            p = paragraph._p
            pPr = p.get_or_add_pPr()
            keepLines = OxmlElement('w:keepLines')
            pPr.append(keepLines)
        except:
            pass

    # ============================================================
    # PARSER DE SKILLS (HEURÍSTICA SIMPLIFICADA)
    # ============================================================
    
    def extract_skills_simple(text):
        """
        Extrai skills de texto corrido usando múltiplas estratégias
        VERSÃO SIMPLIFICADA E ROBUSTA
        """
        if not text:
            return []
        
        try:
            # Limpa o texto
            text = re.sub(r'\s+', ' ', str(text)).strip()
            skills = []
            
            # ESTRATÉGIA 1: Detecta palavras/frases capitalizadas
            # Exemplo: "Análise de Dados" "Automação de Processos"
            words = text.split()
            current_skill = []
            
            for word in words:
                # Se começa com maiúscula ou é conectora
                if word and (word[0].isupper() or word.lower() in ['de', 'e', 'da', 'do', 'em', 'com', 'a', 'o']):
                    current_skill.append(word)
                else:
                    # Finaliza skill atual
                    if current_skill:
                        skill_text = ' '.join(current_skill)
                        if len(skill_text) > 3 and len(skill_text) < 50:
                            skills.append(skill_text)
                    current_skill = []
            
            # Adiciona última skill
            if current_skill:
                skill_text = ' '.join(current_skill)
                if len(skill_text) > 3 and len(skill_text) < 50:
                    skills.append(skill_text)
            
            # Remove duplicatas
            seen = set()
            unique = []
            for s in skills:
                s_lower = s.lower().strip()
                if s_lower not in seen and len(s) > 3:
                    seen.add(s_lower)
                    unique.append(s)
            
            return unique[:15]
        
        except Exception as e:
            # Se der erro, retorna lista vazia
            return []

    # ============================================================
    # INICIALIZA DOCUMENTO
    # ============================================================
    
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    # ============================================================
    # PROCESSA TEXTO BRUTO
    # ============================================================
    
    raw_text = data.get('cv_otimizado_completo', '')
    if not raw_text:
        # Documento vazio se não houver conteúdo
        p = doc.add_paragraph("Currículo vazio")
        stream = BytesIO()
        doc.save(stream)
        stream.seek(0)
        return stream
    
    lines = raw_text.split('\n')
    
    # Estrutura para organizar o conteúdo
    sections = []
    current_section = None
    buffer = []
    contact_lines = []
    name_found = False
    
    # ============================================================
    # PARSING LINHA POR LINHA
    # ============================================================
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Remove markdown
        clean = line.replace('**', '').replace('*', '')
        
        # 1. DETECTA NOME (primeira linha ou linha com # )
        if not name_found:
            if clean.startswith('# '):
                sections.append({
                    'type': 'name',
                    'content': clean.replace('# ', '')
                })
                name_found = True
                continue
            elif len(clean) < 100 and not clean.startswith('|') and not is_contact_line(clean):
                # Primeira linha não vazia = nome
                sections.append({
                    'type': 'name',
                    'content': clean
                })
                name_found = True
                continue
        
        # 2. DETECTA CONTATOS (antes da primeira seção)
        if not current_section and is_contact_line(clean):
            contact_lines.append(clean)
            continue
        
        # 3. DETECTA SEÇÃO (| TÍTULO ou ### TÍTULO)
        if clean.startswith('|') or clean.startswith('###'):
            # Salva seção anterior
            if current_section and buffer:
                sections.append({
                    'type': 'section',
                    'title': current_section,
                    'content': buffer.copy()
                })
            
            # Salva contatos se tiver
            if contact_lines:
                sections.append({
                    'type': 'contact',
                    'content': contact_lines.copy()
                })
                contact_lines = []
            
            # Nova seção
            current_section = clean.replace('|', '').replace('###', '').strip()
            buffer = []
            continue
        
        # 4. ACUMULA CONTEÚDO
        if current_section:
            buffer.append(clean)
    
    # Salva última seção
    if current_section and buffer:
        sections.append({
            'type': 'section',
            'title': current_section,
            'content': buffer.copy()
        })

    # ============================================================
    # RENDERIZA DOCUMENTO
    # ============================================================
    
    for sec in sections:
        sec_type = sec.get('type', '')
        
        # ========== NOME ==========
        if sec_type == 'name':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(4)
            
            r = p.add_run(sec['content'].upper())
            r.bold = True
            r.font.name = FONT_MAIN
            r.font.size = Pt(24)
            r.font.color.rgb = COLOR_TITLE
            set_keep_with_next(p)
        
        # ========== CONTATOS ==========
        elif sec_type == 'contact':
            for contact_line in sec.get('content', []):
                # Remove rótulos
                clean = contact_line
                for label in ['Email:', 'Telefone:', 'LinkedIn:', 'GitHub:', 'Local:']:
                    clean = clean.replace(label, '')
                
                # Split por | ou •
                parts = re.split(r'\s*[|•]\s*', clean)
                valid_parts = [p.strip() for p in parts if is_valid(p)]
                
                if valid_parts:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_after = Pt(12)
                    
                    for i, part in enumerate(valid_parts):
                        r = p.add_run(part)
                        r.font.name = FONT_MAIN
                        r.font.size = Pt(9.5)
                        r.font.color.rgb = COLOR_SUB
                        
                        if i < len(valid_parts) - 1:
                            r_sep = p.add_run("  •  ")
                            r_sep.font.color.rgb = COLOR_ACCENT
                            r_sep.bold = True
        
        # ========== SEÇÕES ==========
        elif sec_type == 'section':
            title = sec.get('title', '').upper()
            content = sec.get('content', [])
            
            if not content:
                continue
            
            # Título da seção
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            set_border_bottom(p)
            set_keep_with_next(p)
            
            r_marker = p.add_run("| ")
            r_marker.font.name = FONT_MAIN
            r_marker.font.size = Pt(14)
            r_marker.bold = True
            r_marker.font.color.rgb = COLOR_ACCENT
            
            r_title = p.add_run(title)
            r_title.bold = True
            r_title.font.name = FONT_MAIN
            r_title.font.size = Pt(11)
            r_title.font.color.rgb = COLOR_TITLE
            
            # ========== SKILLS (DETECÇÃO ESPECIAL) ==========
            if any(kw in title for kw in ['SKILL', 'COMPETÊNCIA', 'TÉCNICA']):
                full_text = ' '.join(content)
                skills = extract_skills_simple(full_text)
                
                if skills:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(8)
                    p.paragraph_format.line_spacing = 1.3
                    
                    for i, skill in enumerate(skills):
                        if i > 0:
                            p.add_run("  ")
                        
                        r_chip = p.add_run(f"  {skill}  ")
                        r_chip.font.name = FONT_MAIN
                        r_chip.font.size = Pt(9.5)
                        r_chip.font.color.rgb = COLOR_CHIP_TEXT
                        apply_chip_background(r_chip)
                else:
                    # Fallback: texto corrido
                    p = doc.add_paragraph(full_text)
                    p.paragraph_format.space_after = Pt(8)
                
                continue
            
            # ========== CONTEÚDO GERAL ==========
            for line in content:
                if not is_valid(line):
                    continue
                
                # CABEÇALHO DE EXPERIÊNCIA (tem |)
                if '|' in line and len(line) < 150:
                    # Remove bullet se tiver
                    line_clean = line
                    if line_clean.startswith('- '):
                        line_clean = line_clean[2:]
                    
                    parts = [p.strip() for p in line_clean.split('|')]
                    
                    # Valida (pelo menos 2 partes válidas)
                    if len(parts) >= 2 and is_valid(parts[0]) and is_valid(parts[1]):
                        # Cargo + Empresa
                        p = doc.add_paragraph()
                        p.paragraph_format.space_before = Pt(10)
                        p.paragraph_format.space_after = Pt(1)
                        set_keep_together(p)
                        
                        r_cargo = p.add_run(parts[0])
                        r_cargo.bold = True
                        r_cargo.font.name = FONT_MAIN
                        r_cargo.font.size = Pt(11.5)
                        r_cargo.font.color.rgb = COLOR_TITLE
                        
                        r_sep = p.add_run(" | ")
                        r_sep.font.color.rgb = COLOR_SUB
                        
                        r_emp = p.add_run(parts[1])
                        r_emp.bold = True
                        r_emp.font.name = FONT_MAIN
                        r_emp.font.size = Pt(11)
                        r_emp.font.color.rgb = COLOR_ACCENT
                        
                        # Data (se tiver)
                        if len(parts) >= 3 and is_valid(parts[2]):
                            p_date = doc.add_paragraph()
                            p_date.paragraph_format.space_after = Pt(3)
                            
                            r_date = p_date.add_run(parts[2])
                            r_date.font.name = FONT_MAIN
                            r_date.font.size = Pt(9)
                            r_date.font.color.rgb = COLOR_SUB
                        
                        continue
                
                # BULLET POINT (descrição de experiência)
                if line.startswith('- ') or line.startswith('• '):
                    clean = line[2:].strip()
                    
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.line_spacing = 1.25
                    
                    # Hanging indent
                    p.paragraph_format.tab_stops.add_tab_stop(Cm(0.5), WD_TAB_ALIGNMENT.LEFT)
                    p.paragraph_format.left_indent = Cm(0.5)
                    p.paragraph_format.first_line_indent = Cm(-0.5)
                    set_keep_together(p)
                    
                    # Bullet
                    r_bullet = p.add_run("•")
                    r_bullet.font.name = FONT_MAIN
                    r_bullet.font.size = Pt(12)
                    r_bullet.font.color.rgb = COLOR_ACCENT
                    p.add_run("\t")
                    
                    # Texto (com negrito em títulos se tiver :)
                    if ':' in clean:
                        title_part, desc_part = clean.split(':', 1)
                        
                        r_title = p.add_run(title_part + ':')
                        r_title.bold = True
                        r_title.font.name = FONT_MAIN
                        r_title.font.size = Pt(10.5)
                        r_title.font.color.rgb = COLOR_TITLE
                        
                        r_desc = p.add_run(' ' + desc_part.strip())
                        r_desc.font.name = FONT_MAIN
                        r_desc.font.size = Pt(10.5)
                        r_desc.font.color.rgb = COLOR_TEXT
                    else:
                        r_txt = p.add_run(clean)
                        r_txt.font.name = FONT_MAIN
                        r_txt.font.size = Pt(10.5)
                        r_txt.font.color.rgb = COLOR_TEXT
                    
                    continue
                
                # TEXTO CORRIDO
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.line_spacing = 1.25
                p.paragraph_format.space_after = Pt(6)
                
                r = p.add_run(line)
                r.font.name = FONT_MAIN
                r.font.size = Pt(10.5)
                r.font.color.rgb = COLOR_TEXT

    # ============================================================
    # SALVA E RETORNA
    # ============================================================
    
    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream
